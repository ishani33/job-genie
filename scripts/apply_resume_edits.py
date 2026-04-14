#!/usr/bin/env python3
"""
Apply bullet edits to a resume .docx template using python-docx.

Preserves all formatting (fonts, sizes, bold, spacing, indentation, columns).
Only the text content of matched bullet paragraphs is changed.

Usage:
    python3 apply_resume_edits.py <template_path> <output_path> <edits_json_path>

edits_json is a list of objects:
    { "section": "Phyllo", "action": "modify"|"add", "original": "...", "editedValue": "..." }
"""

import sys
import json
import re
import copy

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(2)


def normalize(text: str) -> str:
    """Collapse whitespace and lowercase for fuzzy matching."""
    return re.sub(r"\s+", " ", text.strip()).lower()


def para_full_text(para) -> str:
    """Concatenate all run text in a paragraph."""
    return "".join(r.text for r in para.runs)


def set_para_text(para, new_text: str) -> None:
    """
    Replace paragraph text while preserving the formatting of the first run.

    Strategy: set the first run's text to the full new string, then clear all
    subsequent runs. This keeps the paragraph's font, size, bold, italic,
    spacing, indentation, and style object untouched.
    """
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def looks_like_section_header(para) -> bool:
    """
    Heuristic: a section header is short, non-empty, and either:
    - has a Heading style, OR
    - all non-whitespace runs are bold
    """
    text = para_full_text(para).strip()
    if not text or len(text) > 80:
        return False
    if para.style and para.style.name.lower().startswith("heading"):
        return True
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if runs_with_text and all(r.bold for r in runs_with_text):
        return True
    return False


def find_section_range(paragraphs: list, section_name: str):
    """
    Return (start_idx, end_idx) — the slice of paragraphs that belongs to
    the named section. end_idx is exclusive (next section header or EOF).
    Returns (-1, -1) if the section header is not found.
    """
    norm_section = normalize(section_name)
    start = -1

    for i, para in enumerate(paragraphs):
        if normalize(para_full_text(para)) == norm_section:
            start = i
            break
    if start < 0:
        # Try partial / contains match as fallback
        for i, para in enumerate(paragraphs):
            if norm_section in normalize(para_full_text(para)):
                start = i
                break
    if start < 0:
        return (-1, -1)

    # Find where the section ends
    end = len(paragraphs)
    for i in range(start + 1, len(paragraphs)):
        if looks_like_section_header(paragraphs[i]):
            end = i
            break

    return (start, end)


def last_content_para_idx(paragraphs: list, start: int, end: int) -> int:
    """Return the index of the last non-empty paragraph in [start, end)."""
    idx = start
    for i in range(start, end):
        if para_full_text(paragraphs[i]).strip():
            idx = i
    return idx


def apply_edits(template_path: str, output_path: str, edits: list) -> None:
    doc = Document(template_path)

    for edit in edits:
        action = edit.get("action", "keep")
        original = edit.get("original", "")
        new_text = edit.get("editedValue") or edit.get("suggested", "")
        section = edit.get("section", "")

        if action == "keep" or not new_text:
            continue

        paragraphs = doc.paragraphs  # re-fetch after any insertions

        if action == "modify" and original:
            norm_orig = normalize(original)
            for para in paragraphs:
                if normalize(para_full_text(para)) == norm_orig:
                    set_para_text(para, new_text)
                    break
            # Fallback: try substring match if exact match fails
            else:
                best = None
                best_len = 0
                for para in paragraphs:
                    t = normalize(para_full_text(para))
                    if t and t in norm_orig and len(t) > best_len:
                        best = para
                        best_len = len(t)
                if best is not None:
                    set_para_text(best, new_text)

        elif action == "add" and section:
            start, end = find_section_range(paragraphs, section)
            if start < 0:
                # Section not found — append at end of doc as fallback
                doc.add_paragraph(new_text)
                continue

            insert_after_idx = last_content_para_idx(paragraphs, start + 1, end)
            ref_para = paragraphs[insert_after_idx]

            # Deep-copy the reference paragraph's XML node and insert after it
            new_elem = copy.deepcopy(ref_para._element)
            ref_para._element.addnext(new_elem)

            # Find the newly inserted paragraph object and set its text
            paragraphs = doc.paragraphs
            for para in paragraphs:
                if para._element is new_elem:
                    set_para_text(para, new_text)
                    break

    doc.save(output_path)
    print(f"OK: saved to {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print(
            "Usage: python3 apply_resume_edits.py "
            "<template_path> <output_path> <edits_json_path>",
            file=sys.stderr,
        )
        sys.exit(1)

    _, template_path, output_path, edits_json_path = sys.argv

    with open(edits_json_path, "r", encoding="utf-8") as f:
        edits = json.load(f)

    apply_edits(template_path, output_path, edits)
